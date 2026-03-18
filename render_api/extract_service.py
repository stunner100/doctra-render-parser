from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Literal, TypedDict, cast

FileKind = Literal["pdf", "docx", "image"]
Profile = Literal[
    "basic_pdf",
    "enhanced_pdf",
    "paddleocr_vl",
    "docx_structured",
    "image_ocr",
]

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
RAW_PAGE_SPLIT_PATTERN = re.compile(r"\f+")


class CandidateMetrics(TypedDict, total=False):
    tableCount: int
    formulaCount: int
    chartCount: int


class CandidatePage(TypedDict, total=False):
    index: int
    text: str
    chars: int
    tableCount: int
    formulaCount: int
    source: str


class ExtractionCandidate(TypedDict):
    backend: Literal["doctra"]
    kind: FileKind
    parser: Profile
    text: str
    charCount: int
    pageCount: int
    pages: list[CandidatePage]
    warnings: list[str]
    metrics: CandidateMetrics


def load_enhanced_pdf_parser_cls():
    from doctra.parsers.enhanced_pdf_parser import EnhancedPDFParser

    return EnhancedPDFParser


def load_paddleocr_vl_parser_cls():
    from doctra.parsers.paddleocr_vl_parser import PaddleOCRVLPDFParser

    return PaddleOCRVLPDFParser


def load_structured_docx_parser_cls():
    from doctra.parsers.structured_docx_parser import StructuredDOCXParser

    return StructuredDOCXParser


def detect_file_kind(file_name: str, content_type: str | None) -> FileKind:
    normalized_content_type = str(content_type or "").split(";")[0].strip().lower()
    suffix = Path(file_name or "").suffix.lower()

    if normalized_content_type.startswith("image/"):
        return "image"
    if normalized_content_type == PDF_MIME or suffix == ".pdf":
        return "pdf"
    if normalized_content_type == DOCX_MIME or suffix == ".docx":
        return "docx"

    raise ValueError("Unsupported file type. Upload PDF, DOCX, or image.")


def default_profile_for_kind(kind: FileKind) -> Profile:
    if kind == "pdf":
        return "basic_pdf"
    if kind == "docx":
        return "docx_structured"
    return "image_ocr"


def normalize_requested_profile(kind: FileKind, profile: str | None) -> Profile:
    requested = str(profile or "").strip().lower() or default_profile_for_kind(kind)
    normalized = cast(Profile, requested)

    allowed_profiles = {
        "pdf": {"basic_pdf", "enhanced_pdf", "paddleocr_vl"},
        "docx": {"docx_structured"},
        "image": {"image_ocr"},
    }
    if normalized not in allowed_profiles[kind]:
        raise ValueError(f"Unsupported extraction profile '{requested}' for {kind}.")
    return normalized


def normalize_extracted_text(value: str) -> str:
    text = str(value or "")
    text = text.replace("\r\n", "\n")
    text = re.sub(r"!\[[^\]]*\]\(([^)]+)\)", " ", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*>\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^\|?[-:\s|]+\|?$", "", text, flags=re.MULTILINE)
    text = text.replace("|", " ")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def read_first_markdown(root_dir: Path) -> str:
    markdown_files = sorted(root_dir.rglob("*.md"))
    for markdown_path in markdown_files:
        raw = markdown_path.read_text(encoding="utf-8", errors="ignore")
        cleaned = normalize_extracted_text(raw)
        if cleaned:
            return cleaned
    return ""


def build_candidate_pages(value: str, *, target_chars_per_page: int = 2600) -> list[CandidatePage]:
    raw_text = str(value or "").replace("\r\n", "\n")
    if not raw_text.strip():
        return []

    split_pages = [
        normalize_extracted_text(page)
        for page in RAW_PAGE_SPLIT_PATTERN.split(raw_text)
        if normalize_extracted_text(page)
    ]
    if split_pages:
        return [
            {
                "index": index,
                "text": page_text,
                "chars": len(page_text),
                "tableCount": 0,
                "formulaCount": 0,
                "source": "doctra",
            }
            for index, page_text in enumerate(split_pages)
        ]

    text = normalize_extracted_text(raw_text)
    if not text:
        return []

    paragraphs = [chunk.strip() for chunk in re.split(r"\n{2,}", text) if chunk.strip()]
    if not paragraphs:
        return [{
            "index": 0,
            "text": text,
            "chars": len(text),
            "tableCount": 0,
            "formulaCount": 0,
            "source": "doctra",
        }]

    pages: list[CandidatePage] = []
    current = ""
    for paragraph in paragraphs:
        candidate = paragraph if not current else f"{current}\n\n{paragraph}"
        if current and len(candidate) > target_chars_per_page:
            pages.append({
                "index": len(pages),
                "text": current,
                "chars": len(current),
                "tableCount": 0,
                "formulaCount": 0,
                "source": "doctra",
            })
            current = paragraph
            continue
        current = candidate

    if current:
        pages.append({
            "index": len(pages),
            "text": current,
            "chars": len(current),
            "tableCount": 0,
            "formulaCount": 0,
            "source": "doctra",
        })

    return pages


def build_candidate(
    *,
    kind: FileKind,
    parser: Profile,
    text: str,
    warnings: list[str] | None = None,
    metrics: CandidateMetrics | None = None,
) -> ExtractionCandidate:
    normalized_text = normalize_extracted_text(text)
    pages = build_candidate_pages(text)
    if not pages and normalized_text:
        pages = [{
            "index": 0,
            "text": normalized_text,
            "chars": len(normalized_text),
            "tableCount": 0,
            "formulaCount": 0,
            "source": "doctra",
        }]

    return {
        "backend": "doctra",
        "kind": kind,
        "parser": parser,
        "text": normalized_text,
        "charCount": len(normalized_text),
        "pageCount": max(len(pages), 1 if normalized_text else 0),
        "pages": pages,
        "warnings": list(warnings or []),
        "metrics": dict(metrics or {"tableCount": 0, "formulaCount": 0, "chartCount": 0}),
    }


def read_parser_markdown_output(output_root: Path) -> str:
    expected = output_root / "document.md"
    if expected.exists():
        return normalize_extracted_text(expected.read_text(encoding="utf-8", errors="ignore"))
    return read_first_markdown(output_root)


def parse_pdf_basic(pdf_path: Path, *, max_pages: int | None = None, dpi: int = 220) -> str:
    def read_pdf_text_direct(path: Path) -> str:
        import fitz  # PyMuPDF

        chunks: list[str] = []
        with fitz.open(path) as doc:
            for index, page in enumerate(doc):
                if max_pages is not None and index >= max_pages:
                    break
                chunks.append(page.get_text("text") or "")
        return "\n\n".join(chunks)

    def ocr_pdf_with_doctra(path: Path) -> str:
        from doctra.engines.ocr.api import ocr_image
        from doctra.utils.pdf_io import render_pdf_to_images

        lines: list[str] = []
        rendered_pages = render_pdf_to_images(str(path), dpi=dpi)
        for index, page_item in enumerate(rendered_pages):
            if max_pages is not None and index >= max_pages:
                break
            pil_image = page_item[0].convert("RGB")
            page_text = ocr_image(pil_image, lang="eng", psm=4, oem=3)
            if page_text:
                lines.append(page_text)
        return "\n\n".join(lines)

    direct_text = normalize_extracted_text(read_pdf_text_direct(pdf_path))
    if len(direct_text) >= 250:
        return direct_text

    ocr_text = normalize_extracted_text(ocr_pdf_with_doctra(pdf_path))
    if len(ocr_text) > len(direct_text):
        return ocr_text
    return direct_text


def select_pdf_parser(profile: Profile):
    if profile == "enhanced_pdf":
        parser_cls = load_enhanced_pdf_parser_cls()
        return parser_cls(
            use_image_restoration=True,
            merge_split_tables=True,
        )
    if profile == "paddleocr_vl":
        parser_cls = load_paddleocr_vl_parser_cls()
        return parser_cls(
            use_image_restoration=True,
            use_chart_recognition=True,
            merge_split_tables=True,
        )
    raise ValueError(f"Unsupported PDF parser profile '{profile}'.")


def parse_pdf(pdf_path: Path, stem: str, profile: Profile, *, max_pages: int | None = None) -> tuple[str, list[str], CandidateMetrics]:
    if profile == "basic_pdf":
        return parse_pdf_basic(pdf_path, max_pages=max_pages), [], {
            "tableCount": 0,
            "formulaCount": 0,
            "chartCount": 0,
        }

    warnings: list[str] = []
    output_root = Path("outputs") / stem
    try:
        parser = select_pdf_parser(profile)
        parser.parse(str(pdf_path))
        parsed_text = read_parser_markdown_output(output_root)
        if parsed_text:
            return parsed_text, warnings, {
                "tableCount": 0,
                "formulaCount": 0,
                "chartCount": 0,
            }
        warnings.append(f"{profile}_empty_output")
    except Exception:
        warnings.append(f"{profile}_failed")

    fallback_text = parse_pdf_basic(pdf_path, max_pages=max_pages)
    return fallback_text, warnings, {
        "tableCount": 0,
        "formulaCount": 0,
        "chartCount": 0,
    }


def parse_docx(docx_path: Path, stem: str, profile: Profile) -> tuple[str, list[str], CandidateMetrics]:
    if profile != "docx_structured":
        raise ValueError(f"Unsupported DOCX parser profile '{profile}'.")

    def read_docx_text_direct(path: Path) -> str:
        from docx import Document

        doc = Document(path)
        lines: list[str] = []

        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [(cell.text or "").strip() for cell in row.cells]
                row_text = " | ".join([cell for cell in cells if cell])
                if row_text:
                    lines.append(row_text)

        return "\n".join(lines)

    direct_text = normalize_extracted_text(read_docx_text_direct(docx_path))
    if len(direct_text) >= 120:
        return direct_text, [], {
            "tableCount": 0,
            "formulaCount": 0,
            "chartCount": 0,
        }

    parser_cls = load_structured_docx_parser_cls()
    parser = parser_cls(
        extract_images=False,
        table_detection=True,
        export_excel=False,
    )
    parser.parse(str(docx_path))
    output_root = Path("outputs") / stem
    parsed_text = read_parser_markdown_output(output_root)
    return parsed_text, [], {
        "tableCount": 0,
        "formulaCount": 0,
        "chartCount": 0,
    }


def parse_image(image_path: Path, profile: Profile) -> tuple[str, list[str], CandidateMetrics]:
    if profile != "image_ocr":
        raise ValueError(f"Unsupported image parser profile '{profile}'.")

    from PIL import Image
    from doctra.engines.ocr.api import ocr_image

    with Image.open(image_path) as img:
        normalized = img.convert("RGB")
    text = normalize_extracted_text(ocr_image(normalized, lang="eng", psm=4, oem=3))
    return text, [], {
        "tableCount": 0,
        "formulaCount": 0,
        "chartCount": 0,
    }


def extract_candidate(
    file_path: Path,
    file_name: str,
    content_type: str | None,
    profile: str | None = None,
    max_pages: int | None = None,
) -> ExtractionCandidate:
    kind = detect_file_kind(file_name=file_name, content_type=content_type)
    normalized_profile = normalize_requested_profile(kind, profile)
    stem = file_path.stem
    output_root = Path("outputs") / stem

    try:
        if kind == "pdf":
            text, warnings, metrics = parse_pdf(
                file_path,
                stem,
                normalized_profile,
                max_pages=max_pages,
            )
        elif kind == "docx":
            text, warnings, metrics = parse_docx(file_path, stem, normalized_profile)
        else:
            text, warnings, metrics = parse_image(file_path, normalized_profile)
    finally:
        if output_root.exists():
            shutil.rmtree(output_root, ignore_errors=True)

    return build_candidate(
        kind=kind,
        parser=normalized_profile,
        text=text,
        warnings=warnings,
        metrics=metrics,
    )


def extract_text(file_path: Path, file_name: str, content_type: str | None) -> tuple[str, FileKind]:
    candidate = extract_candidate(
        file_path=file_path,
        file_name=file_name,
        content_type=content_type,
    )
    return candidate["text"], candidate["kind"]
